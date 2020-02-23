# python-docx module for documents with .docx files
# use pip install python-docx to install the library
# use import docx to use in python once the module is installed.

# Document object represents the entire document
# The Document object contains a list of Paragraph objects for the paragraphs in the document
# Each of these Paragraph objects contains a list of one or more Run objects
# A Run object is a contiguous run of text with the same style. A new Run object is needed whenever the text style changes.

# docObj = docx.Document(file), Docbj.paragraph, docObj.paragrpah.text, docObj.paragraph.run, docObj.paragraph.run.text, docObj.save()
# Document(NO_FILENAME) : creates a new document.
# docObj.add_paragraph('string') : creates and adds the contents to teh file as new paragraph and returns the ref obj of teh paragraph.
# paraObj.add_run('string') : adds content to the paragraph
# docObj.save(newfilename)

# add_paragraph() and add_run() accept an optional second argument that is a string of the Paragraph or Run object’s style. For example:
# doc.add_paragraph('Hello world!', 'Title')   : This line adds a paragraph with the text Hello world! in the Title style.

# The text in a Word document is more than just a string. It has font, size, color, and other styling information associated with it.
# A style in Word is a collection of these attributes.
# There are three types of styles for Word doc: Paragraph styles, character styles, linked styles
# Paragraph Style: can be apied to Paragraph objects, character styles: can be applied to Run objects, linked styles can be applied to both kinds of objects.
# You can give both Paragraph and Run objects styles by setting their style attribute to a string value. This string should be the name of any existing style. ex: 'Normal', 'BodyText'
# set the style to None if no style needs to be associated with Paragraph/run objects
# When setting the style attribute, do not use spaces in the style name. For example use 'SubtleEmphasis' instead of 'Subtle Emphasis'

# When using a linked style for a Run object, you will need to add 'Char' to the end of its name.
# For example, to set the Quote linked style for a Paragraph object, use paragraphObj.style = 'Quote', but for a Run object, you would use runObj.style = 'QuoteChar'

# Runs can be further styled using text attributes. True/False/None   . True-> Enable, False-> Disable, None-> whatever is there set to
# ex: doc.paragraphs[1].runs[1].underline = True

# Calling add_heading() adds a paragraph with one of the heading styles.  ex: doc.add_heading('Content', 0)   . heading Styles(Integer) : 0, 1,2, 3,4
# add_heading() also returns the paraObj.   docObj.add_heading('Content', int_style) or parObj = docObj.add_heading('Content', int_style)

# Adding Line and Page Breaks : add_break() on Run object can be used to create a line/page break instead of creating a new paragraph
# pass docx.text.WD_BREAK.PAGE as argument to the add_break() to create a page break
# doc.paragraphs[0].runs[0].add_break(docx.text.run.WD_BREAK.PAGE)

# docObj.add_picture() method can add an image to the end of the document.
# it needs Image.ext as first argument, and optional 2 arguments for width and height :
# width=docx.shared.Inches(), height=docx.shared.Cm()



import docx

#=====================================Reading Document=============================================
doc = docx.Document('demo.docx')
len(doc.paragraphs)                            # 7
print(len(doc.paragraphs))
for i in range(len(doc.paragraphs)):
    print(f"{i} is {doc.paragraphs[i].text}")

doc.paragraphs[0].text                         # 'Document Title'
doc.paragraphs[1].text                         # 'A plain paragraph with some bold and some italic'
len(doc.paragraphs[1].runs)                    # 4
doc.paragraphs[1].runs[0].text                 # 'A plain paragraph with some '
doc.paragraphs[1].runs[1].text                 # 'bold'
doc.paragraphs[1].runs[2].text                 #' and some '
doc.paragraphs[1].runs[3].text                 # 'italic'


#=================================== Geeting Full text from a .docx file=======================
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)

    return '\n'.join(fullText) # fullText.append(' ' + para.text)  # return '\n\n'.join(fullText)

# ============================ create a new word document =================================
doc = docx.Document()
doc.add_paragraph('Hello world!')

paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')

paraObj1.add_run(' This text is being added to the second paragraph.')

doc.save('multipleParagraphs.docx')

"""
Hello World!
This is a second paragraph. This text is being added to the second paragraph.
This text is being added to the second paragraph.
"""

# =================================Adding Headings =================================
# Calling add_heading() adds a paragraph with one of the heading styles.

doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')

# ============================= Adding Line and Page Breaks=========================
doc = docx.Document()
doc.add_paragraph('This is on the first page!')

doc.paragraphs[0].runs[0].add_break()
doc.paragraphs[0].runs[1].add_break(docx.text.run.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')

doc.save('twoPage.docx')

# ============== Ading Picture ===============================
doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))

# it can take both width=docx.shared.Inches() : Inches,        height=docx.shared.Cm() : Centimeter
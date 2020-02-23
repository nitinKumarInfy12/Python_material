#! python3

"""Take a list of names and make a docx with custom invites for each one."""

import docx

with open('guests.txt') as f:
    names = f.readlines()
    document = docx.Document()

print(names)

for name in names:
    name = name.strip()
    document.add_paragraph('It would be a pleasure to have the company of', style='Title')
    document.add_paragraph('\t\t'+ name, style='Quote')
    document.add_paragraph('at 11010 Memory Lane on the Evening of', style='MacroText')
    prObj1 = document.add_paragraph('April 1st', style='IntenseQuote')
    prObj1.add_run("at 7 o'clock").bold=True
    document.add_page_break()

    document.save(name+'_invites.docx')

print("File has been created and saved as 'invites.docx'")

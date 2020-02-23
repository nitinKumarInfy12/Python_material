import docx

doc = docx.Document()
doc.add_paragraph('This is on the first page!')

for i in range(len(doc.paragraphs)):
    print(len(doc.paragraphs[i].runs))
    for j in range(len(doc.paragraphs[i].runs)):
        print(doc.paragraphs[i].runs[j].text)
doc.paragraphs[0].runs[0].add_break()
doc.paragraphs[0].runs[1].add_break(docx.text.run.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')

doc.save('twoPage.docx')